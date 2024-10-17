import os
import fitz  # PyMuPDF for handling PDF files
import pdfplumber  # For extracting tables from PDFs
import csv  # For saving tables as CSV files
from docx.oxml.ns import qn  # Used for namespacing in DOCX processing
import sys
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
from loaders.file_loader import AbstractFileLoader
from loaders.pdf_loader import PDFLoader
from loaders.ppt_loader import PPTLoader
from loaders.docx_loader import DOCXLoader

def clean_text(text):
    """
    Cleans extracted text by removing unwanted characters like \t and \n, and strips any leading/trailing whitespace.
    Args:
        text (str): The text to clean.
    Returns:
        str: Cleaned text.
    """
    return text.replace("\n", " ").replace("\t", " ").strip()

class DataExtractor:
    def __init__(self, loader):
        """
        Initializes the DataExtractor with a specific file loader instance.
        Args:
            loader (PDFLoader | DOCXLoader | PPTLoader): The loader instance capable of loading a specific file format.
        """
        self.loader = loader

    def extract_text(self):
        """
        Extracts text from a loaded file using the appropriate loader.
        Differentiates extraction logic based on the file type.
        Returns:
            list | dict: Text data extracted from the file, formatted according to file type.
        """
        if isinstance(self.loader, PDFLoader):
            return self._extract_pdf_text(self.loader.filepath)  # Special handling for PDF files directly from the path

        loaded_file = self.loader.open_file(self.loader.filepath)  # Load file for DOCX or PPT

        if isinstance(self.loader, DOCXLoader):
            return self._extract_docx_text(loaded_file)
        elif isinstance(self.loader, PPTLoader):
            return self._extract_pptx_text(loaded_file)

    def _extract_pdf_text(self, pdf_path):
        """
        Extracts text from a PDF file, merging text blocks intelligently to maintain logical content structure.
        Args:
            pdf_path (str): The file path to the PDF document.
        Returns:
            list: List of dictionaries with page numbers and content for each page.
        """
        doc = fitz.open(pdf_path)  # Open the PDF document using PyMuPDF
        text_data = []

        for page_num in range(len(doc)):
            page = doc.load_page(page_num)  # Load each page individually
            blocks = page.get_text("dict")["blocks"]  # Extract text in 'dict' format to get structured blocks
            page_content = []
            current_line = ""
            current_style = None  # Style tracking variable

            for block in blocks:
                if "lines" in block:
                    for line in block["lines"]:
                        line_text = ""
                        line_style = None

                        for span in line["spans"]:
                            font_size = span["size"]
                            text = span["text"].strip()
                            style = "Heading" if font_size > 14 else "normal"

                            line_text += " " + text if line_text else text
                            line_style = line_style or style

                        # Continuously merge text or start new line based on style consistency
                        if current_line and line_style == current_style:
                            current_line += " " + line_text
                        else:
                            if current_line:  # Finish the current line and start a new one
                                page_content.append({"text": current_line.strip(), "style": current_style})
                            current_line = line_text
                            current_style = line_style

            # Ensure the last line of the page is added
            if current_line:
                page_content.append({"text": current_line.strip(), "style": current_style})

            text_data.append({"page_number": page_num + 1, "content": page_content})

        return text_data

    def _extract_docx_text(self, doc):
        """
        Extracts text from a DOCX file and returns a list of dictionaries,
        each containing the text and its associated style if it has one.
        Args:
            doc (Document): The loaded DOCX file object from python-docx.

        Returns:
            list: A list of dictionaries with keys 'text' and 'style' representing each paragraph's content and style name.
        """
        # List comprehension to iterate over all paragraphs in the document,
        # clean the text, and collect text and style name if the paragraph is not empty.
        return [
            {"text": clean_text(paragraph.text), "style": paragraph.style.name if paragraph.style else "Normal"}
            for paragraph in doc.paragraphs if paragraph.text.strip()
        ]

    def _extract_pptx_text(self, presentation):
        """
        Extracts text from a PPTX file and compiles it into a structured list, considering text frames within shapes on each slide.
        This function also handles basic text styling by identifying headings based on bolding and font size.

        Args:
            presentation (Presentation): The loaded PPTX file object from python-pptx.

        Returns:
            list: A list of dictionaries, each containing slide number and content,
                which includes cleaned text and styles.
        """
        text_data = []

        # Loop through each slide in the presentation
        for slide_num, slide in enumerate(presentation.slides):
            slide_content = []

            # Check each shape in the slide; focus on those with text frames
            for shape in slide.shapes:
                if shape.has_text_frame:
                    for paragraph in shape.text_frame.paragraphs:
                        paragraph_text = ""
                        style = "normal"  # Default style

                        # Concatenate all runs in the paragraph to form the full text
                        for run in paragraph.runs:
                            paragraph_text += run.text

                            # Determine style by checking if the text is bold or font size is significantly large
                            if run.font.bold or (run.font.size and run.font.size > 200000):
                                style = "Heading"

                        # Clean text and filter out any paragraph that consists only of whitespace
                        cleaned_text = clean_text(paragraph_text)
                        if cleaned_text:
                            slide_content.append({
                                "text": cleaned_text,
                                "style": style
                            })

            # Only include slides that contain content to avoid empty entries
            if slide_content:
                text_data.append({
                    "slide_number": slide_num + 1,
                    "content": slide_content
                })

        return text_data

    def extract_links(self):
        """
        Extracts hyperlinks from the currently loaded file using the appropriate loader.
        Determines the file type from the loader instance and calls the corresponding method
        to handle hyperlink extraction specific to each file type.

        Returns:
            list: A list of dictionaries, each containing metadata about the hyperlinks found.
        """
        loaded_file = self.loader.open_file(self.loader.filepath)

        if isinstance(self.loader, PDFLoader):
            return self._extract_pdf_links(loaded_file)
        elif isinstance(self.loader, DOCXLoader):
            return self._extract_docx_links(loaded_file)
        elif isinstance(self.loader, PPTLoader):
            return self._extract_pptx_links(loaded_file)

    def _extract_pdf_links(self, pdf_reader):
        """
        Extracts hyperlinks from a PDF file using annotations, which are often used to store hyperlink data.
        Args:
            pdf_reader (fitz.Document): The loaded PDF document.

        Returns:
            list: A list of dictionaries where each dictionary contains the page number and the hyperlink URL.
        """
        links_data = []
        for page_num, page in enumerate(pdf_reader.pages):
            annotations = page.get('/Annots')  # Retrieve annotations from the page
            if annotations:
                for annotation in annotations:
                    # Attempt to retrieve the hyperlink URI from the annotation
                    uri = annotation.get_object().get('/A').get('/URI')
                    if uri:
                        links_data.append({
                            "page_number": page_num + 1,  # Page numbers are indexed from 1 for user clarity
                            "link": uri
                        })
        return links_data

    def _extract_docx_links(self, doc):
        """
        Extracts hyperlinks from a DOCX file, focusing only on the URLs and ensuring no duplicates are stored.
        Args:
            doc (Document): The loaded DOCX file object from python-docx.

        Returns:
            list: A list of dictionaries, each containing a unique hyperlink ('link') from the document.
        """
        links_data = []  # Initialize the list to hold link data.
        extracted_links = set()  # Use a set to track already extracted links to avoid duplicates.

        # Iterate over all relationships in the document. Each 'rel' represents a link, image, or other external reference.
        for rel in doc.part.rels.values():
            # Check if the relationship is a hyperlink.
            if "hyperlink" in rel.reltype:
                # Ensure the hyperlink has not already been extracted.
                if rel.target_ref not in extracted_links:
                    extracted_links.add(rel.target_ref)  # Add the hyperlink to the set of extracted links.
                    links_data.append({"link": rel.target_ref})  # Store the hyperlink in the list.

        return links_data

    def _extract_pptx_links(self, presentation):
        """
        Extracts hyperlinks from a PPTX file, capturing both the linked text and the hyperlink address.
        Args:
            presentation (Presentation): The loaded PPTX file object from python-pptx.

        Returns:
            list: A list of dictionaries, each containing the slide number, linked text, and the hyperlink URL.
        """
        links_data = []  # Initialize the list to hold link data.

        # Iterate over all slides in the presentation.
        for slide_num, slide in enumerate(presentation.slides):
            # Iterate over all shapes in the slide that have a text frame.
            for shape in slide.shapes:
                if shape.has_text_frame:
                    # Check each paragraph in the text frame.
                    for paragraph in shape.text_frame.paragraphs:
                        linked_text = ""
                        link = None

                        # Check each run in the paragraph for hyperlinks.
                        for run in paragraph.runs:
                            if run.hyperlink and run.hyperlink.address:
                                # Assign the first hyperlink address found and accumulate the text associated with the hyperlink.
                                link = link or run.hyperlink.address
                                linked_text += run.text

                        # If a hyperlink was found and has associated text, store it, ensuring no duplicate entries.
                        if link and linked_text and not any(d['link'] == link and d['linked_text'] == clean_text(linked_text) for d in links_data):
                            links_data.append({
                                "slide_number": slide_num + 1,  # 1-based index for user clarity.
                                "linked_text": clean_text(linked_text),  # Cleaned text to ensure consistency.
                                "link": link
                            })

        return links_data

    def extract_images(self):
        """
        Extract images based on the file type of the loaded document. Determines the type of loader and
        delegates to the appropriate image extraction method.
        """
        loaded_file = self.loader.load_file(self.loader.filepath)  # Load the file using the appropriate loader
        if isinstance(self.loader, PDFLoader):
            return self._extract_pdf_images(self.loader.filepath)  # Extract images from PDF
        elif isinstance(self.loader, DOCXLoader):
            return self._extract_docx_images(loaded_file)  # Extract images from DOCX
        elif isinstance(self.loader, PPTLoader):
            return self._extract_pptx_images(loaded_file)  # Extract images from PPTX

    def _extract_pdf_images(self, pdf_path):
        """
        Extracts all images from a PDF file and saves them locally.
        Args:
            pdf_path (str): The file path to the PDF document.

        Returns:
            list: A list of dictionaries containing details about each extracted image.
        """
        images_data = []
        doc = fitz.open(pdf_path)  # Open the PDF document using PyMuPDF
        pdf_images_folder = os.path.join("output", "images", "pdf")  # Define the directory to store images
        os.makedirs(pdf_images_folder, exist_ok=True)  # Ensure the directory exists

        for page_num, page in enumerate(doc.pages()):  # Iterate through each page in the PDF
            for image_index, image in enumerate(page.get_images(full=True)):  # Get all images from the page
                xref = image[0]  # Reference number for the image
                base_image = doc.extract_image(xref)  # Extract the image using its reference
                image_filename = f"pdf_image_{page_num+1}_{image_index+1}.{base_image['ext']}"  # Create a filename
                image_path = os.path.join(pdf_images_folder, image_filename)  # Create a full path for the image

                with open(image_path, "wb") as image_file:  # Write the image file to disk
                    image_file.write(base_image["image"])  # Save the image data

                # Append image details to the list
                images_data.append({
                    "page_number": page_num + 1,
                    "image_filename": image_filename,
                    "image_format": base_image["ext"],
                    "image_path": image_path
                })

        return images_data

    def _extract_docx_images(self, doc):
        """
        Extract images from a DOCX file and save them to a specified directory.
        Args:
            doc (Document): The loaded DOCX document object.

        Returns:
            list: A list of dictionaries, each containing metadata about the extracted images.
        """
        images_data = []
        docx_images_folder = os.path.join("output", "images", "docx")
        os.makedirs(docx_images_folder, exist_ok=True)  # Ensure the output directory exists

        # Iterate through all inline shapes in the document that are images
        for i, shape in enumerate(doc.inline_shapes):
            # Access the binary data of the image
            image_part = doc.part.related_parts[shape._inline.graphic.graphicData.pic.blipFill.blip.embed]
            image_filename = f"docx_image_{i+1}.{image_part.content_type.split('/')[-1]}"  # Construct filename
            image_path = os.path.join(docx_images_folder, image_filename)  # Construct file path

            # Write the image file to the disk
            with open(image_path, "wb") as image_file:
                image_file.write(image_part.blob)

            # Append image details to the list for later use or reference
            images_data.append({
                "image_filename": image_filename,
                "image_format": image_part.content_type.split('/')[-1],
                "image_path": image_path
            })
        return images_data

    def _extract_pptx_images(self, presentation):
        """
        Extract images from a PPTX file, specifically from slides that contain image shapes.
        Args:
            presentation (Presentation): The loaded PPTX file object.

        Returns:
            list: A list of dictionaries detailing the images extracted from each slide.
        """
        images_data = []
        pptx_images_folder = os.path.join("output", "images", "pptx")
        os.makedirs(pptx_images_folder, exist_ok=True)  # Ensure the output directory exists

        # Iterate through each slide and its shapes to find images
        for slide_num, slide in enumerate(presentation.slides):
            for shape in slide.shapes:
                if shape.shape_type == 13:  # Picture type in PowerPoint
                    image = shape.image
                    image_filename = f"pptx_image_{slide_num+1}_{shape.shape_id}.{image.ext}"  # Construct filename
                    image_path = os.path.join(pptx_images_folder, image_filename)  # Construct file path

                    # Write the image file to the disk
                    with open(image_path, "wb") as image_file:
                        image_file.write(image.blob)

                    # Append image details to the list for later use or reference
                    images_data.append({
                        "slide_number": slide_num + 1,
                        "image_filename": image_filename,
                        "image_format": image.ext,
                        "image_path": image_path
                    })
        return images_data

    def extract_tables(self):
        """
        Extract tables based on the file type of the loaded document. Determines the type of loader and
        delegates to the appropriate table extraction method.
        """
        loaded_file = self.loader.load_file(self.loader.filepath)  # Load the file using the appropriate loader
        if isinstance(self.loader, PDFLoader):
            return self._extract_pdf_tables(self.loader.filepath)  # Extract tables from PDF
        elif isinstance(self.loader, DOCXLoader):
            return self._extract_docx_tables(loaded_file)  # Extract tables from DOCX
        elif isinstance(self.loader, PPTLoader):
            return self._extract_pptx_tables(loaded_file)  # Extract tables from PPTX

    def _extract_pdf_tables(self, pdf_path):
        """
        Extracts tables from a PDF file, processes them page by page, and saves them as CSV files.
        Each table extracted is saved into a separate CSV file named distinctly by page and table index.

        Args:
            pdf_path (str): The file path to the PDF document.

        Returns:
            list: A list of dictionaries containing metadata about the extracted tables and their CSV file paths.
        """
        tables_data = []  # List to store metadata about the extracted tables
        pdf_tables_folder = os.path.join("output", "tables", "pdf")  # Define the directory to store CSV files
        os.makedirs(pdf_tables_folder, exist_ok=True)  # Ensure the directory exists

        with pdfplumber.open(pdf_path) as pdf:  # Open the PDF with pdfplumber
            for page_num, page in enumerate(pdf.pages):  # Iterate through each page in the PDF
                tables = page.extract_tables()  # Extract all tables found on the current page
                for table_index, table in enumerate(tables):  # Iterate through each table
                    csv_filename = f"pdf_table_{page_num+1}_{table_index+1}.csv"  # Create a unique filename for the CSV
                    csv_path = os.path.join(pdf_tables_folder, csv_filename)  # Create the full path for the CSV file

                    # Write the table data to a CSV file
                    with open(csv_path, 'w', newline='', encoding='utf-8') as csvfile:
                        writer = csv.writer(csvfile)
                        writer.writerows(table)  # Write each row of the table to the CSV file

                    # Store metadata about the table in the list
                    tables_data.append({
                        "page_number": page_num + 1,  # Page number (1-indexed for readability)
                        "table_index": table_index + 1,  # Table index (1-indexed for readability)
                        "csv_filename": csv_filename,
                        "csv_path": csv_path
                    })

        return tables_data

    def _extract_docx_tables(self, doc):
        """
        Extracts tables from a DOCX file and saves them as CSV files in a specified directory.
        Each table is saved into a separate CSV file named uniquely based on its index in the document.
        
        Args:
            doc (Document): The loaded DOCX document object from python-docx.

        Returns:
            list: A list of dictionaries containing metadata about the extracted tables and their CSV file paths.
        """
        tables_data = []  # Initialize a list to hold metadata about each extracted table
        docx_tables_folder = os.path.join("output", "tables", "docx")  # Define the directory to store CSV files
        os.makedirs(docx_tables_folder, exist_ok=True)  # Ensure the directory exists

        # Iterate over each table in the document
        for table_index, table in enumerate(doc.tables):
            csv_filename = f"docx_table_{table_index+1}.csv"  # Construct a unique filename for the CSV
            csv_path = os.path.join(docx_tables_folder, csv_filename)  # Create the full path for the CSV file

            # Open a new CSV file and write the table data
            with open(csv_path, 'w', newline='', encoding='utf-8') as csvfile:
                writer = csv.writer(csvfile)
                writer.writerows([[cell.text for cell in row.cells] for row in table.rows])  # Convert table rows to CSV

            # Append metadata about the table to the list
            tables_data.append({
                "table_index": table_index + 1,  # Index is 1-based for user clarity
                "csv_filename": csv_filename,
                "csv_path": csv_path
            })
        return tables_data

    def _extract_pptx_tables(self, presentation):
        """
        Extracts tables from a PPTX file and saves them as CSV files in a specified directory.
        Each table is saved into a separate CSV file named uniquely based on its slide number and shape ID.

        Args:
            presentation (Presentation): The loaded PPTX presentation object from python-pptx.

        Returns:
            list: A list of dictionaries detailing the tables extracted from each slide, including CSV file paths.
        """
        tables_data = []  # Initialize a list to hold metadata about each extracted table
        pptx_tables_folder = os.path.join("output", "tables", "pptx")  # Define the directory to store CSV files
        os.makedirs(pptx_tables_folder, exist_ok=True)  # Ensure the directory exists

        # Iterate through each slide in the presentation
        for slide_num, slide in enumerate(presentation.slides):
            # Check each shape on the slide for a table
            for shape in slide.shapes:
                if shape.has_table:
                    table = shape.table  # Get the table object
                    csv_filename = f"pptx_table_{slide_num+1}_{shape.shape_id}.csv"  # Construct a unique filename for the CSV
                    csv_path = os.path.join(pptx_tables_folder, csv_filename)  # Create the full path for the CSV file

                    # Open a new CSV file and write the table data
                    with open(csv_path, 'w', newline='', encoding='utf-8') as csvfile:
                        writer = csv.writer(csvfile)
                        writer.writerows([[cell.text for cell in row.cells] for row in table.rows])  # Convert table rows to CSV

                    # Append metadata about the table to the list
                    tables_data.append({
                        "slide_number": slide_num + 1,  # Slide number is 1-based for user clarity
                        "csv_filename": csv_filename,
                        "csv_path": csv_path
                    })
        return tables_data