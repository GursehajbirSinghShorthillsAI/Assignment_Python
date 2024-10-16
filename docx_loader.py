import os
import fitz  # PyMuPDF for PDFs
import csv
from PIL import Image
from docx import Document  # For DOCX files
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pptx
from abc import ABC, abstractmethod
import csv
from PyPDF2 import PdfReader
import json


class DOCXLoader:
    def __init__(self, file_path):
        self.file_path = file_path
        self.doc = Document(file_path)

    def extract_detailed_metadata(self):
        metadata = []
        for para in self.doc.paragraphs:
            para_metadata = []
            for run in para.runs:
                para_metadata.append({
                    "text": run.text,
                    "font": run.font.name,
                    "size": run.font.size.pt if run.font.size else None,  # Get font size in points
                    "bold": run.bold,
                    "italic": run.italic,
                    "underline": run.underline
                })
            metadata.append({
                "paragraph_text": para.text,
                "run_metadata": para_metadata
            })
        return metadata

    def extract_text(self):
        # Extract all the text from the DOCX document
        text = "\n".join([para.text for para in self.doc.paragraphs])
        return text

    def extract_links(self):
        # Extract links from the DOCX document
        links = []
        for rel in self.doc.part.rels.values():
            if "hyperlink" in rel.target_ref:
                links.append(rel.target_ref)
        return links
    

    def extract_tables(self):
        # Extract tables from DOCX document
        tables = []
        for table in self.doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        return tables